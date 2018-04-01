from docx.shared import Pt

from generation import Project

class Stage7(object):

    def __init__(self, erd):
        self.erd = erd

    def build(self, document):
        header = document.add_paragraph()
        header.add_run(u'7. Definicje encji i związków').font.size = Pt(Project.HEADER_SIZE)
        header.add_run().add_break()

        entities_paragraph = document.add_paragraph()
        entities_paragraph.add_run('7.1 Encje').font.size = Pt(Project.SECONDAR_HEADER_SIZE)
        entities_paragraph.add_run().add_break()

        for entity in self.erd.entities:
            entity_paragraph = document.add_paragraph()
            entity_paragraph.add_run('ENC/' + '{0:03}'.format(entity.id) + ' ' + entity.name_singular.upper()).bold = True
            entity_paragraph.add_run().add_break()
            entity_paragraph.add_run('\tSemantyka encji:\n').italic = True
            entity_paragraph.add_run('\tWykaz atrybutów:').italic = True

            table = document.add_table(rows=len(entity.attributes) + 1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].paragraphs[0].add_run('Nazwa atrybutu').bold = True
            hdr_cells[1].paragraphs[0].add_run('Opis atrybutu').bold = True
            hdr_cells[2].paragraphs[0].add_run('Typ').bold = True
            hdr_cells[3].paragraphs[0].add_run('OBL(+), OPC(-)').bold = True

            counter = 1
            for attribute in entity.attributes:
                row = table.rows[counter].cells
                row[0].text = attribute.name
                row[2].text = repr(attribute.type)

                counter += 1

            entities_paragraph2 = document.add_paragraph()
            entities_paragraph2.add_run('Klucze kandydujące:').italic = True
            entities_paragraph2.add_run(' ').add_break()
            entities_paragraph2.add_run('Charakter encji:').italic = True
            entities_paragraph2.add_run(' ').add_break()